from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
from werkzeug.utils import secure_filename
import uuid
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pdf2docx import Converter
import logging
import re

app = Flask(__name__)
CORS(app)

app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def analyze_pdf(pdf_path):
    """Analyze PDF content type using pdfplumber"""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_text = 0
            total_tables = 0
            page_count = len(pdf.pages)
            
            for page in pdf.pages[:3]:  # Sample first 3 pages
                text = page.extract_text()
                if text:
                    total_text += len(text)
                tables = page.extract_tables()
                total_tables += len(tables)
            
            text_ratio = total_text / (page_count * 500)  # Avg chars per page
            has_tables = total_tables > 0
            
            if text_ratio > 0.8:
                return "text_perfect"
            elif has_tables:
                return "table_heavy"
            else:
                return "image_layout"
    except:
        return "image_layout"

def perfect_text_conversion(pdf_path, docx_path):
    """Perfect text extraction with formatting preservation"""
    logger.info("✨ Using PERFECT text converter")
    
    doc = Document()
    
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            logger.info(f"Processing page {page_num + 1}")
            
            # Extract text blocks (preserves layout)
            text_blocks = page.extract_words(x_tolerance=3, y_tolerance=3, layout=True)
            
            if text_blocks:
                # Group words into lines by Y position
                lines = {}
                for word in text_blocks:
                    y_pos = round(word['top'], 1)
                    if y_pos not in lines:
                        lines[y_pos] = []
                    lines[y_pos].append(word['text'])
                
                # Add lines as paragraphs
                for y_pos in sorted(lines.keys()):
                    line_text = ' '.join(lines[y_pos])
                    if line_text.strip():
                        p = doc.add_paragraph(line_text.strip())
                        p.paragraph_format.space_after = Inches(0.08)
                        p.paragraph_format.line_spacing = 1.2
            
            # Extract and format tables
            tables = page.extract_tables()
            for table_idx, table in enumerate(tables):
                if table and len(table) > 1 and len(table[0]) > 1:
                    rows = len(table)
                    cols = len(table[0])
                    
                    t = doc.add_table(rows=rows, cols=cols)
                    t.style = 'Table Grid'
                    
                    for i, row in enumerate(table):
                        for j, cell in enumerate(row or []):
                            if i < len(t.rows) and j < len(t.rows[i].cells):
                                t.rows[i].cells[j].text = str(cell).strip() if cell else ''
                    
                    # Add space after table
                    doc.add_paragraph()
    
    doc.save(docx_path)
    logger.info("✅ Perfect text conversion complete!")

@app.route('/')
def home():
    return jsonify({
        'status': 'Perfect PDF to Word (Text + Images) ✅',
        'engines': ['Advanced pdfplumber (text)', 'pdf2docx (images)']
    })

@app.route('/api/health')
def health():
    return jsonify({'status': 'healthy'})

@app.route('/api/convert', methods=['POST'])
def convert_pdf():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '' or not allowed_file(file.filename):
            return jsonify({'error': 'Invalid PDF file'}), 400
        
        # Save uploaded PDF
        unique_id = str(uuid.uuid4())[:8]
        filename = f"{unique_id}_{secure_filename(file.filename)}"
        pdf_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(pdf_path)
        
        docx_path = pdf_path.replace('.pdf', '.docx')
        
        # Auto-detect best conversion method
        pdf_type = analyze_pdf(pdf_path)
        logger.info(f"📊 PDF Type: {pdf_type}")
        
        if pdf_type == "text_perfect":
            perfect_text_conversion(pdf_path, docx_path)
        else:
            logger.info("🖼️ Using layout-preserving pdf2docx")
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
        
        # Cleanup original PDF
        os.remove(pdf_path)
        
        return jsonify({
            'success': True,
            'pdf_type': pdf_type,
            'download_url': f'/api/download/{os.path.basename(docx_path)}',
            'filename': os.path.basename(docx_path)
        })
        
    except Exception as e:
        logger.error(f"❌ Error: {str(e)}")
        return jsonify({'error': f'Conversion failed: {str(e)}'}), 500

@app.route('/api/download/<filename>')
def download_file(filename):
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    if os.path.exists(filepath):
        return send_file(filepath, 
                        as_attachment=True,
                        download_name=filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    return jsonify({'error': 'File not found or expired'}), 404

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
